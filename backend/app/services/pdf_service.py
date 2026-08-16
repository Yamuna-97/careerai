"""
app/services/pdf_service.py
────────────────────────────
PDF export service using ReportLab (pure Python, no system dependencies).

Generates a clean, professional PDF resume from the resume model data.
The layout can be customized per template in the future.
"""

import io
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from app.models.resume import Resume


# ── Color palette ─────────────────────────────────────────────────────────────
PRIMARY_COLOR = colors.HexColor("#3525cd")
DARK_TEXT = colors.HexColor("#1a1a2e")
MUTED_TEXT = colors.HexColor("#6b7280")
DIVIDER = colors.HexColor("#e5e7eb")


def generate_resume_pdf(resume: Resume) -> bytes:
    """
    Generate a PDF resume from a Resume model instance.
    Returns the PDF as bytes (to be sent as a file download).
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Name & Contact ─────────────────────────────────────────────────
    name_style = ParagraphStyle("Name", fontSize=22, textColor=PRIMARY_COLOR,
                                 fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
    contact_style = ParagraphStyle("Contact", fontSize=9, textColor=MUTED_TEXT,
                                    alignment=TA_CENTER, spaceAfter=2)
    section_style = ParagraphStyle("Section", fontSize=12, fontName="Helvetica-Bold",
                                    textColor=PRIMARY_COLOR, spaceBefore=12, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontSize=9, textColor=DARK_TEXT,
                                 leading=14, spaceAfter=3)
    sub_style = ParagraphStyle("Sub", fontSize=9, textColor=MUTED_TEXT, spaceAfter=2)

    # Header
    story.append(Paragraph(resume.full_name or "Your Name", name_style))

    contact_parts = []
    if resume.email: contact_parts.append(resume.email)
    if resume.phone: contact_parts.append(resume.phone)
    if resume.location: contact_parts.append(resume.location)
    if resume.linkedin: contact_parts.append(resume.linkedin)
    if resume.github: contact_parts.append(resume.github)

    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1.5, color=PRIMARY_COLOR, spaceAfter=8))

    # ── Summary ────────────────────────────────────────────────────────
    if resume.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        story.append(Paragraph(resume.summary, body_style))

    # ── Education ─────────────────────────────────────────────────────
    if resume.education:
        story.append(Paragraph("EDUCATION", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        for edu in resume.education:
            date_range = f"{edu.start_date or ''} – {edu.end_date or 'Present'}".strip(" –")
            story.append(Paragraph(
                f"<b>{edu.degree or ''} {edu.field_of_study or ''}</b>", body_style
            ))
            story.append(Paragraph(
                f"{edu.institution} | {date_range}" + (f" | {edu.grade}" if edu.grade else ""),
                sub_style
            ))
            if edu.description:
                story.append(Paragraph(edu.description, body_style))
            story.append(Spacer(1, 4))

    # ── Experience ────────────────────────────────────────────────────
    if resume.experience:
        story.append(Paragraph("EXPERIENCE", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        for exp in resume.experience:
            end = "Present" if exp.currently_working else (exp.end_date or "")
            date_range = f"{exp.start_date or ''} – {end}".strip(" –")
            story.append(Paragraph(
                f"<b>{exp.position}</b> — {exp.company}", body_style
            ))
            story.append(Paragraph(
                f"{exp.location or ''} | {date_range}".strip(" |"), sub_style
            ))
            if exp.description:
                story.append(Paragraph(exp.description, body_style))
            story.append(Spacer(1, 4))

    # ── Projects ──────────────────────────────────────────────────────
    if resume.projects:
        story.append(Paragraph("PROJECTS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        for proj in resume.projects:
            story.append(Paragraph(f"<b>{proj.name}</b>", body_style))
            if proj.technologies:
                story.append(Paragraph(f"Technologies: {proj.technologies}", sub_style))
            if proj.description:
                story.append(Paragraph(proj.description, body_style))
            links = []
            if proj.github_url: links.append(f"GitHub: {proj.github_url}")
            if proj.live_url: links.append(f"Live: {proj.live_url}")
            if links:
                story.append(Paragraph(" | ".join(links), sub_style))
            story.append(Spacer(1, 4))

    # ── Skills ────────────────────────────────────────────────────────
    if resume.skills:
        story.append(Paragraph("SKILLS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        # Group by category
        by_category: dict[str, list[str]] = {}
        for skill in resume.skills:
            by_category.setdefault(skill.category, []).append(skill.name)
        for category, skill_names in by_category.items():
            story.append(Paragraph(
                f"<b>{category}:</b> {', '.join(skill_names)}", body_style
            ))

    # ── Certifications ────────────────────────────────────────────────
    if resume.certifications:
        story.append(Paragraph("CERTIFICATIONS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        for cert in resume.certifications:
            story.append(Paragraph(
                f"<b>{cert.name}</b> — {cert.issuer or ''} {cert.issue_date or ''}".strip(" —"),
                body_style
            ))

    # ── Achievements ──────────────────────────────────────────────────
    if resume.achievements:
        story.append(Paragraph("ACHIEVEMENTS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=DIVIDER))
        story.append(Spacer(1, 4))
        for ach in resume.achievements:
            story.append(Paragraph(
                f"<b>{ach.title}</b>" + (f" — {ach.organization}" if ach.organization else ""),
                body_style
            ))
            if ach.description:
                story.append(Paragraph(ach.description, body_style))

    # ── Build PDF ─────────────────────────────────────────────────────
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ── Text Extraction Helpers for Uploaded Files ─────────────────────────────────

def extract_text_from_pdf(content_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes."""
    if not content_bytes:
        return ""

    # 1. Try pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                pages_text.append(t.strip())
        if pages_text:
            return "\n\n".join(pages_text)
    except Exception as e:
        print(f"[pdf_service] pypdf extraction warning: {e}")

    # 2. Raw regex fallback for unencrypted PDFs
    try:
        import re
        raw_str = content_bytes.decode("latin1", errors="ignore")
        matches = re.findall(r'\((.*?)\)\s*Tj', raw_str)
        if matches:
            return " ".join(matches)
    except Exception:
        pass

    return ""


def extract_text_from_docx(content_bytes: bytes) -> str:
    """Extract plain text from Word (.docx) file bytes."""
    if not content_bytes:
        return ""

    # 1. Try python-docx
    try:
        import docx
        doc = docx.Document(io.BytesIO(content_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paragraphs.append(cell.text.strip())
        if paragraphs:
            return "\n".join(paragraphs)
    except Exception as e:
        print(f"[pdf_service] docx extraction warning: {e}")

    # 2. XML fallback via zipfile parsing
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(content_bytes)) as z:
            if "word/document.xml" in z.namelist():
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                text_nodes = tree.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                text_list = [node.text for node in text_nodes if node.text]
                if text_list:
                    return "\n".join(text_list)
    except Exception as e:
        print(f"[pdf_service] docx zip XML extraction failed: {e}")

    return ""
