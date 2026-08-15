"""
app/services/resume_ai_service.py
──────────────────────────────────
Gemini AI integration for resume creation, improvement, extraction,
ATS optimization, scoring, and chat-based editing.

All Gemini calls return strictly validated JSON.
Safety guarantee: AI is explicitly instructed to NEVER invent
experience, companies, degrees, skills, or certifications.
"""

import httpx
import json
import io
from typing import List, Optional, Any
from app.core.config import settings

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent"


# ── Internal Gemini Helper ──────────────────────────────────────────────────────
def _call_gemini(prompt: str, json_mode: bool = True) -> str:
    """POST to Gemini API and return raw text response."""
    if not settings.GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not configured in .env")

    headers = {"Content-Type": "application/json"}
    params = {"key": settings.GEMINI_API_KEY}
    payload: dict = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 8192,
        },
    }
    if json_mode:
        payload["generationConfig"]["responseMimeType"] = "application/json"

    response = httpx.post(
        GEMINI_API_URL, headers=headers, params=params, json=payload, timeout=45.0
    )
    response.raise_for_status()
    data = response.json()
    candidates = data.get("candidates", [])
    if not candidates:
        raise ValueError("Gemini returned no candidates.")
    parts = candidates[0].get("content", {}).get("parts", [])
    if not parts:
        raise ValueError("Gemini returned empty parts.")
    return parts[0].get("text", "")


def _parse_json(raw: str) -> dict:
    """Parse JSON response from Gemini, stripping markdown fences if present."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        cleaned = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    return json.loads(cleaned)


# ── Resume Schema Template ─────────────────────────────────────────────────────
RESUME_SCHEMA = {
    "personal": {
        "fullName": "", "title": "", "email": "", "phone": "",
        "location": "", "linkedin": "", "github": "", "portfolio": "", "profileImage": ""
    },
    "summary": "",
    "education": [],
    "experience": [],
    "projects": [],
    "skills": [],
    "certifications": [],
    "achievements": [],
    "languages": []
}

SAFETY_PREAMBLE = """
CRITICAL SAFETY RULE: You MUST NOT invent, fabricate, or add any of the following:
- Companies, employers, or organizations
- Job titles or roles the person did not mention
- Degrees, universities, or academic credentials
- Skills, technologies, or tools not mentioned
- Certifications not mentioned
- Projects, repositories, or work not mentioned
- Achievements, awards, or competitions not mentioned
- Any dates, numbers, or statistics not explicitly provided

You may ONLY: rewrite, rephrase, improve clarity, fix grammar, use stronger action verbs,
improve formatting, and make descriptions more concise or impactful.
If a section is empty or not mentioned, leave it empty in your output.
"""


# ── 1. Extract Resume Data from Text ──────────────────────────────────────────
def extract_resume_data(text: str) -> dict:
    """
    Parse raw resume text (from PDF/DOCX or paste) into structured JSON
    matching CareerAI's resume schema.
    """
    prompt = f"""
You are a resume parser. Extract structured information from the following resume text.

{SAFETY_PREAMBLE}

Return ONLY valid JSON matching this exact schema:
{{
  "personal": {{
    "fullName": "string",
    "title": "string (professional title/role)",
    "email": "string",
    "phone": "string",
    "location": "string",
    "linkedin": "string (URL or handle)",
    "github": "string (URL or handle)",
    "portfolio": "string (URL)",
    "profileImage": ""
  }},
  "summary": "string (professional summary paragraph)",
  "education": [
    {{
      "id": "1",
      "institution": "string",
      "degree": "string",
      "fieldOfStudy": "string",
      "startDate": "string",
      "endDate": "string",
      "grade": "string",
      "description": "string"
    }}
  ],
  "experience": [
    {{
      "id": "1",
      "company": "string",
      "position": "string",
      "location": "string",
      "startDate": "string",
      "endDate": "string",
      "currentlyWorking": false,
      "description": "string"
    }}
  ],
  "projects": [
    {{
      "id": "1",
      "name": "string",
      "description": "string",
      "technologies": "string (comma-separated)",
      "githubUrl": "string",
      "liveUrl": "string",
      "startDate": "string",
      "endDate": "string"
    }}
  ],
  "skills": [
    {{
      "id": "1",
      "name": "string",
      "category": "Programming Languages|Frameworks|Databases|Machine Learning|Tools|Cloud|Other"
    }}
  ],
  "certifications": [
    {{
      "id": "1",
      "name": "string",
      "issuer": "string",
      "issueDate": "string",
      "credentialUrl": "string",
      "description": "string"
    }}
  ],
  "achievements": [
    {{
      "id": "1",
      "title": "string",
      "organization": "string",
      "date": "string",
      "description": "string"
    }}
  ],
  "languages": []
}}

Resume text to parse:
---
{text}
---

Return ONLY the JSON, no explanation, no markdown.
"""
    raw = _call_gemini(prompt, json_mode=True)
    data = _parse_json(raw)
    # Add sequential IDs to arrays
    for section in ["education", "experience", "projects", "skills", "certifications", "achievements"]:
        items = data.get(section, [])
        for i, item in enumerate(items):
            item["id"] = str(i + 1)
    return data


# ── 2. Generate Polished Resume ────────────────────────────────────────────────
def generate_resume(
    data: dict,
    target_role: str = "",
    job_description: str = "",
    tone: str = "Professional"
) -> dict:
    """
    Take structured resume data and return an AI-improved version.
    Gemini will improve wording, grammar, impact, and tailor to target role.
    It MUST NOT add any fabricated information.
    """
    context = f"Target Role: {target_role}" if target_role else ""
    job_context = f"\nJob Description for tailoring:\n{job_description[:2000]}" if job_description else ""

    prompt = f"""
You are an expert resume writer. Improve the following resume data to be more professional,
impactful, and ATS-friendly. Tone: {tone}.
{context}
{job_context}

{SAFETY_PREAMBLE}

Improvement guidelines:
- Use strong action verbs (Led, Developed, Architected, Optimized, Implemented, Delivered)
- Make bullet points concise and impact-focused
- Improve the professional summary to be compelling and role-focused
- Ensure skill names use industry-standard terminology
- Fix any grammar or phrasing issues
- Make project descriptions highlight technical impact
- Do NOT add skills, experience, or certifications not in the original data

Input resume data (JSON):
{json.dumps(data, indent=2)}

Return ONLY the improved resume data as valid JSON in the exact same schema as the input.
Do NOT add new fields. Do NOT remove existing sections that have data.
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw)


# ── 3. Improve Entire Resume ───────────────────────────────────────────────────
def improve_resume(data: dict) -> dict:
    """General improvement pass on the entire resume."""
    return generate_resume(data, target_role="", job_description="", tone="Professional")


# ── 4. Rewrite a Specific Section ─────────────────────────────────────────────
def rewrite_section(
    section: str,
    content: Any,
    instruction: str,
    context: dict
) -> Any:
    """
    Apply a specific instruction to a single resume section.
    Returns only the updated section content.
    """
    prompt = f"""
You are an expert resume editor. Apply the following instruction to this specific resume section.

{SAFETY_PREAMBLE}

Section: {section}
Current content:
{json.dumps(content, indent=2)}

Instruction from user: "{instruction}"

Full resume context (for tone consistency):
Name: {context.get('personal', {}).get('fullName', '')}
Role: {context.get('personal', {}).get('title', '')}

Return ONLY the updated section content as valid JSON in the same structure as the input.
Do not include any explanation or markdown. Just the updated JSON value.
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw) if raw.strip().startswith("{") or raw.strip().startswith("[") else raw.strip()


# ── 5. Generate Professional Summary ──────────────────────────────────────────
def generate_summary(data: dict, target_role: str = "") -> str:
    """Generate a compelling professional summary based on resume data."""
    prompt = f"""
Write a professional resume summary (2-3 sentences, max 60 words) for the following person.

{SAFETY_PREAMBLE}

Target Role: {target_role or data.get('personal', {}).get('title', 'Professional')}
Name: {data.get('personal', {}).get('fullName', '')}
Experience: {json.dumps(data.get('experience', []), indent=2)}
Skills: {json.dumps([s.get('name') for s in data.get('skills', [])], indent=2)}
Education: {json.dumps(data.get('education', []), indent=2)}
Projects: {json.dumps([p.get('name') for p in data.get('projects', [])], indent=2)}

Return ONLY a JSON object: {{"summary": "the generated summary text"}}
"""
    raw = _call_gemini(prompt, json_mode=True)
    result = _parse_json(raw)
    return result.get("summary", "")


# ── 6. Optimize for ATS ────────────────────────────────────────────────────────
def optimize_for_ats(data: dict, job_description: str = "") -> dict:
    """
    Analyze resume against job description for ATS compatibility.
    Returns score, matched keywords, missing keywords, and suggestions.
    Does NOT add fabricated skills.
    """
    prompt = f"""
You are an ATS (Applicant Tracking System) expert. Analyze this resume against the job description.

{SAFETY_PREAMBLE}

Resume data:
{json.dumps(data, indent=2)}

Job description:
{job_description[:3000] if job_description else "No job description provided. Analyze for general ATS best practices."}

Return ONLY valid JSON:
{{
  "ats_score": 0-100,
  "overall_score": 0-100,
  "content_score": 0-100,
  "impact_score": 0-100,
  "readability_score": 0-100,
  "professionalism_score": 0-100,
  "matched_keywords": ["keyword1", "keyword2"],
  "missing_keywords": [
    {{
      "keyword": "Docker",
      "reason": "Job description requires Docker but resume does not mention it",
      "add_if_you_have": true
    }}
  ],
  "weak_sections": ["summary", "experience"],
  "suggestions": [
    {{
      "section": "summary",
      "issue": "Summary is too generic",
      "fix": "Make it role-specific with key technologies",
      "priority": "high"
    }}
  ],
  "improved_resume": null
}}

CRITICAL: For missing_keywords, only list keywords from the job description that are absent from the resume.
Do NOT suggest adding skills the person doesn't have. The add_if_you_have field must always be true (never fabricate).
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw)


# ── 7. Analyze Job Description ────────────────────────────────────────────────
def analyze_job_description(job_description: str) -> dict:
    """Extract key requirements, skills, and keywords from a job description."""
    prompt = f"""
Analyze this job description and extract structured information.

Job Description:
{job_description[:3000]}

Return ONLY valid JSON:
{{
  "job_title": "string",
  "company": "string or empty",
  "required_skills": ["skill1", "skill2"],
  "preferred_skills": ["skill1"],
  "key_responsibilities": ["responsibility1"],
  "experience_required": "string (e.g. 2+ years)",
  "education_required": "string",
  "keywords": ["keyword1", "keyword2"],
  "tech_stack": ["technology1", "technology2"],
  "soft_skills": ["communication", "teamwork"]
}}
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw)


# ── 8. Score Resume ────────────────────────────────────────────────────────────
def score_resume(data: dict, target_role: str = "") -> dict:
    """Score a resume across multiple dimensions."""
    prompt = f"""
Score this resume across multiple professional dimensions.

Target Role: {target_role or data.get('personal', {}).get('title', 'Professional')}

Resume:
{json.dumps(data, indent=2)}

Return ONLY valid JSON:
{{
  "overall_score": 0-100,
  "ats_score": 0-100,
  "content_score": 0-100,
  "impact_score": 0-100,
  "readability_score": 0-100,
  "professionalism_score": 0-100,
  "keyword_match_score": 0-100,
  "summary": "One sentence about the resume's main strength",
  "top_strengths": ["strength1", "strength2", "strength3"],
  "improvement_areas": ["area1", "area2"]
}}
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw)


# ── 9. Generate Improvement Suggestions ───────────────────────────────────────
def generate_resume_suggestions(data: dict) -> list:
    """Return prioritized list of improvement suggestions."""
    prompt = f"""
Review this resume and provide specific, actionable improvement suggestions.

Resume:
{json.dumps(data, indent=2)}

Return ONLY valid JSON as an array:
[
  {{
    "id": "1",
    "section": "summary|experience|projects|skills|achievements|education|certifications",
    "priority": "high|medium|low",
    "issue": "Short description of the problem",
    "suggestion": "Specific actionable suggestion",
    "fix_prompt": "Instruction to give AI to fix this (e.g. 'Rewrite the summary with stronger action verbs')"
  }}
]

Provide 5-8 suggestions. Focus on the most impactful improvements.
"""
    raw = _call_gemini(prompt, json_mode=True)
    result = _parse_json(raw)
    return result if isinstance(result, list) else result.get("suggestions", [])


# ── 10. Chat Edit Resume ───────────────────────────────────────────────────────
def chat_edit_resume(
    message: str,
    resume_data: dict,
    chat_history: List[dict],
    selected_section: Optional[str] = None
) -> dict:
    """
    Process a natural language editing command and return updated resume + AI response.
    Maintains conversation context.
    """
    history_text = ""
    for msg in chat_history[-6:]:  # Keep last 6 messages for context
        role = "User" if msg.get("role") == "user" else "Assistant"
        history_text += f"{role}: {msg.get('content', '')}\n"

    section_context = f"\nFocused section: {selected_section}" if selected_section else ""

    prompt = f"""
You are an expert AI resume editor. The user wants to edit their resume using natural language commands.

{SAFETY_PREAMBLE}

Current Resume:
{json.dumps(resume_data, indent=2)}
{section_context}

Conversation history:
{history_text}

User's current instruction: "{message}"

Understand what the user wants to change and apply ONLY that change to the resume.
If the user asks to edit only one section, modify only that section.
If the user asks to make summary shorter, only modify the summary field.
If the user asks to remove a section, set it to empty array/string.
If the user asks to add more impact to experience, improve only the experience descriptions.

Return ONLY valid JSON:
{{
  "ai_response": "Brief friendly explanation of what you changed (1-2 sentences)",
  "updated_resume": {{ ... complete updated resume data in same schema ... }},
  "changed_sections": ["list of section names that were modified"]
}}

The updated_resume must include ALL fields, not just the changed ones.
"""
    raw = _call_gemini(prompt, json_mode=True)
    return _parse_json(raw)


# ── 11. Extract Text from PDF ─────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except ImportError:
        raise ValueError("pypdf is not installed. Run: pip install pypdf")
    except Exception as e:
        raise ValueError(f"Failed to extract PDF text: {str(e)}")


# ── 12. Extract Text from DOCX ────────────────────────────────────────────────
def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except ImportError:
        raise ValueError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {str(e)}")


# ── 13. Extract Job Search Profile from Resume ───────────────────────────────
def extract_job_search_profile(resume_data: dict) -> dict:
    """
    Analyze structured resume data and extract initial job search preferences to seed a JobSearchProfile.
    """
    if settings.GEMINI_API_KEY and settings.AI_PROVIDER != "none":
        try:
            prompt = f"""
You are an expert career consultant. Analyze the following resume data and extract initial job search settings.

Resume Data:
{json.dumps(resume_data, indent=2)}

Based on this resume, infer:
1. Target Roles (e.g. ['Software Engineer', 'Frontend Developer'] - list 1 to 3 roles, matching their experience and title)
2. Top 8-10 Skills (e.g. ['React', 'JavaScript', 'TailwindCSS'])
3. Relevant search keywords (e.g. ['SaaS', 'design systems', 'FastAPI'])
4. Experience Level: one of "entry", "junior", "mid", "senior", or "any"
5. Current Title: their most recent professional title, or empty string if none
6. Education Level: their highest degree obtained, or empty string
7. Locations: list of cities or "Remote" (infer from location in resume, default to ["Remote"] if none)
8. Preferred Work Modes: list containing "remote", "hybrid", and/or "onsite" (default to ["remote", "hybrid", "onsite"])
9. Employment Types: list containing "full_time", "part_time", "internship", and/or "contract" (default to ["full_time"])
10. Country Code: "in" for India, "us" for USA, "gb" for UK, "ca" for Canada, "au" for Australia, or "in" by default.

Return ONLY a valid JSON object matching this exact schema:
{{
  "target_roles": ["string"],
  "skills": ["string"],
  "keywords": ["string"],
  "experience_level": "entry|junior|mid|senior|any",
  "current_title": "string",
  "education_level": "string",
  "locations": ["string"],
  "work_modes": ["remote"|"hybrid"|"onsite"],
  "employment_types": ["full_time"|"part_time"|"internship"|"contract"],
  "country_code": "in|us|gb|ca|au"
}}
"""
            raw = _call_gemini(prompt, json_mode=True)
            return _parse_json(raw)
        except Exception as e:
            # Fall back to heuristic extraction on failure
            print(f"[!] Error in Gemini extract_job_search_profile: {e}")
            pass

    # Heuristic fallback parsing logic
    personal = resume_data.get("personal", {})
    title = personal.get("title", "") or ""
    
    # Extract skills
    skills_list = [s.get("name") for s in resume_data.get("skills", []) if s.get("name")]
    
    # Target roles
    target_roles = [title] if title else ["Software Engineer"]
    
    # Locations
    loc = personal.get("location", "")
    locations = [loc] if loc else ["Remote"]
    
    # Highest Education Level
    edu_list = resume_data.get("education", [])
    highest_edu = ""
    if edu_list:
        highest_edu = edu_list[0].get("degree", "")
        if edu_list[0].get("fieldOfStudy"):
            highest_edu += f" in {edu_list[0].get('fieldOfStudy')}"

    return {
        "target_roles": target_roles,
        "skills": skills_list[:8],
        "keywords": skills_list[:5],
        "experience_level": "any",
        "current_title": title,
        "education_level": highest_edu,
        "locations": locations,
        "work_modes": ["remote", "hybrid", "onsite"],
        "employment_types": ["full_time"],
        "country_code": "in"
    }

